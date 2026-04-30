# v5/backend/app/services/export_service.py
# 导出服务

from flask import current_app
from ..models.models import db, Novel, Chapter, Character
import os
import uuid


def export_to_docx(novel, options):
    """导出为Word"""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(novel.title, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加简介
        if novel.summary:
            doc.add_paragraph(f"\n{novel.summary}\n")
        
        # 添加大纲（可选）
        if options.get('include_outline') and novel.outline and novel.outline.content:
            doc.add_heading('故事大纲', level=1)
            doc.add_paragraph(novel.outline.content)
        
        # 添加角色介绍（可选）
        if options.get('include_characters'):
            characters = Character.query.filter_by(novel_id=novel.id).all()
            if characters:
                doc.add_heading('角色介绍', level=1)
                for char in characters:
                    doc.add_heading(char.name, level=2)
                    if char.personality:
                        doc.add_paragraph(f"性格：{char.personality}")
                    if char.background:
                        doc.add_paragraph(f"背景：{char.background}")
                    if char.goal:
                        doc.add_paragraph(f"目标：{char.goal}")
        
        # 添加章节
        chapters = Chapter.query.filter_by(novel_id=novel.id)\
            .order_by(Chapter.chapter_no).all()
        
        # 筛选章节范围
        chapter_range = options.get('chapter_range')
        if chapter_range:
            start, end = chapter_range
            chapters = [c for c in chapters if start <= c.chapter_no <= end]
        
        for chapter in chapters:
            if chapter.content:
                doc.add_heading(chapter.title or f'第{chapter.chapter_no}章', level=1)
                doc.add_paragraph(chapter.content)
        
        # 保存文件
        export_folder = current_app.config['EXPORT_FOLDER']
        filename = f"{novel.title}_{uuid.uuid4().hex[:8]}.docx"
        file_path = os.path.join(export_folder, filename)
        doc.save(file_path)
        
        return {
            'filename': filename,
            'download_url': f'/exports/{filename}',
            'message': 'Word导出成功'
        }
    except Exception as e:
        raise


def export_to_txt(novel, options):
    """导出为TXT"""
    try:
        content_parts = []
        
        # 添加标题
        content_parts.append(f"{'='*50}")
        content_parts.append(f"  {novel.title}")
        content_parts.append(f"{'='*50}\n")
        
        # 添加简介
        if novel.summary:
            content_parts.append(f"简介：\n{novel.summary}\n")
        
        # 添加大纲（可选）
        if options.get('include_outline') and novel.outline and novel.outline.content:
            content_parts.append(f"\n{'='*50}")
            content_parts.append("  故事大纲")
            content_parts.append(f"{'='*50}")
            content_parts.append(novel.outline.content)
        
        # 添加角色介绍（可选）
        if options.get('include_characters'):
            characters = Character.query.filter_by(novel_id=novel.id).all()
            if characters:
                content_parts.append(f"\n{'='*50}")
                content_parts.append("  角色介绍")
                content_parts.append(f"{'='*50}")
                for char in characters:
                    content_parts.append(f"\n【{char.name}】")
                    if char.personality:
                        content_parts.append(f"性格：{char.personality}")
                    if char.background:
                        content_parts.append(f"背景：{char.background}")
                    if char.goal:
                        content_parts.append(f"目标：{char.goal}")
        
        # 添加章节
        chapters = Chapter.query.filter_by(novel_id=novel.id)\
            .order_by(Chapter.chapter_no).all()
        
        chapter_range = options.get('chapter_range')
        if chapter_range:
            start, end = chapter_range
            chapters = [c for c in chapters if start <= c.chapter_no <= end]
        
        content_parts.append(f"\n{'='*50}")
        content_parts.append("  正文")
        content_parts.append(f"{'='*50}")
        
        for chapter in chapters:
            if chapter.content:
                content_parts.append(f"\n\n{'─'*30}")
                content_parts.append(f"  {chapter.title or f'第{chapter.chapter_no}章'}")
                content_parts.append(f"{'─'*30}\n")
                content_parts.append(chapter.content)
        
        # 保存文件
        export_folder = current_app.config['EXPORT_FOLDER']
        filename = f"{novel.title}_{uuid.uuid4().hex[:8]}.txt"
        file_path = os.path.join(export_folder, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(content_parts))
        
        return {
            'filename': filename,
            'download_url': f'/exports/{filename}',
            'message': 'TXT导出成功'
        }
    except Exception as e:
        raise


def export_to_epub(novel, options):
    """导出为EPUB"""
    try:
        from ebooklib import epub
        
        book = epub.EpubBook()
        
        # 设置元数据
        book.set_identifier(str(novel.id))
        book.set_title(novel.title)
        book.set_language('zh')
        book.add_author('AI辅助创作')
        
        # 添加样式
        style = '''
        body { font-family: "Microsoft YaHei", sans-serif; line-height: 1.8; padding: 20px; }
        h1 { text-align: center; margin-bottom: 30px; }
        h2 { margin-top: 20px; }
        p { text-indent: 2em; margin-bottom: 10px; }
        '''
        css = epub.EpubItem(
            uid="style",
            file_name="style/default.css",
            media_type="text/css",
            content=style.encode('utf-8')
        )
        book.add_item(css)
        
        # 创建章节
        chapters = Chapter.query.filter_by(novel_id=novel.id)\
            .order_by(Chapter.chapter_no).all()
        
        chapter_range = options.get('chapter_range')
        if chapter_range:
            start, end = chapter_range
            chapters = [c for c in chapters if start <= c.chapter_no <= end]
        
        epub_chapters = []
        toc = []
        spine = ['nav']
        
        for i, chapter in enumerate(chapters):
            if chapter.content:
                # 创建章节
                epub_chapter = epub.EpubHtml(
                    title=chapter.title or f'第{chapter.chapter_no}章',
                    file_name=f'chapter_{chapter.chapter_no}.xhtml',
                    lang='zh'
                )
                
                # 构建HTML内容
                html_content = f'''
                <html>
                <head><link rel="stylesheet" href="style/default.css" type="text/css"/></head>
                <body>
                    <h1>{chapter.title or f'第{chapter.chapter_no}章'}</h1>
                    {"".join(f"<p>{p}</p>" for p in chapter.content.split('\n') if p.strip())}
                </body>
                </html>
                '''
                
                epub_chapter.content = html_content.encode('utf-8')
                epub_chapter.add_item(css)
                
                book.add_item(epub_chapter)
                epub_chapters.append(epub_chapter)
                toc.append(epub_chapter)
                spine.append(epub_chapter)
        
        # 设置目录和书脊
        book.toc = toc
        book.spine = spine
        
        # 添加导航
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        
        # 保存文件
        export_folder = current_app.config['EXPORT_FOLDER']
        filename = f"{novel.title}_{uuid.uuid4().hex[:8]}.epub"
        file_path = os.path.join(export_folder, filename)
        
        epub.write_epub(file_path, book)
        
        return {
            'filename': filename,
            'download_url': f'/exports/{filename}',
            'message': 'EPUB导出成功'
        }
    except Exception as e:
        raise